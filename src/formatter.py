"""
输出格式化器：将转录结果格式化为 Markdown (.md) 和/或 Word (.docx) 文件。

Markdown 输出为纯文本段落，可选择在每段前添加 [HH:MM:SS] 时间戳前缀。
DOCX 输出使用 python-docx 生成带基本样式（字体、字号）的 Word 文档。
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Dict, List


# ============================================================
# 时间戳格式化
# ============================================================

def _format_timestamp(seconds: float) -> str:
    """将秒数转换为 [HH:MM:SS] 格式的时间戳字符串。

    浮点秒数会被四舍五入到最近的整数秒。

    示例:
        >>> _format_timestamp(0.0)
        '[00:00:00]'
        >>> _format_timestamp(3661.5)
        '[01:01:02]'
        >>> _format_timestamp(125.3)
        '[00:02:05]'

    Args:
        seconds: 距离音频起点的秒数（可为浮点数）

    Returns:
        格式为 [HH:MM:SS] 的时间戳字符串
    """
    total_secs = int(round(seconds))
    hours = total_secs // 3600
    minutes = (total_secs % 3600) // 60
    secs = total_secs % 60
    return f"[{hours:02d}:{minutes:02d}:{secs:02d}]"


# ============================================================
# Markdown 输出
# ============================================================

def save_markdown(
    segments: List[Dict[str, Any]],
    output_path: str,
    with_timestamps: bool = True,
) -> str:
    """将转录段落保存为 Markdown 纯文本文件。

    每个段落之间以空行分隔。当 with_timestamps=True 时，
    每个段落以 [HH:MM:SS] 时间戳前缀开头（使用段落的 start 时间）。

    如果 segments 为空，将生成一个包含"未检测到语音内容"提示的空输出文件。

    Args:
        segments: 转录段落列表，每项包含 {'start': float, 'end': float, 'text': str}
        output_path: 输出 .md 文件的完整路径
        with_timestamps: 是否在每段前添加时间戳前缀

    Returns:
        写入的 .md 文件路径

    示例输出::

        [00:00:00] 大家好，欢迎参加今天的会议。

        [00:00:05] 今天我们讨论一下项目进度。

        [00:00:12] 首先请张三介绍一下前端的情况。
    """
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    if not segments:
        # 空转录：生成占位内容
        content = (
            "# 转录结果\n\n"
            "> ⚠️ **未检测到语音内容**\n\n"
            "该音频文件转录为空，可能原因：\n"
            "- 音频为纯音乐或纯器乐，不含人声\n"
            "- 音频全程为静音或环境噪音\n"
            "- 语言与模型支持的语言不匹配\n\n"
            "---\n"
            "*此文件由 Video-Transcriber 自动生成*\n"
        )
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        return output_path

    lines: List[str] = []

    for seg in segments:
        text = seg.get("text", "").strip()
        if not text:
            continue

        if with_timestamps:
            ts = _format_timestamp(seg["start"])
            lines.append(f"{ts} {text}")
        else:
            lines.append(text)

    content = "\n\n".join(lines) + "\n"

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)

    return output_path


# ============================================================
# DOCX 输出
# ============================================================

def save_docx(
    segments: List[Dict[str, Any]],
    output_path: str,
    with_timestamps: bool = True,
) -> str:
    """将转录段落保存为 Word (.docx) 文档。

    使用 python-docx 生成带基本样式的文档：
    - 正文字体：Calibri，字号 11pt
    - 时间戳使用等宽字体 Consolas，颜色为深灰色，便于区分
    - 段落间距：段前 0pt，段后 6pt
    - 行距：1.15 倍行距

    如果 segments 为空，将生成一个包含"未检测到语音内容"提示的空输出文件。

    Args:
        segments: 转录段落列表，每项包含 {'start': float, 'end': float, 'text': str}
        output_path: 输出 .docx 文件的完整路径
        with_timestamps: 是否在每段前添加时间戳前缀

    Returns:
        写入的 .docx 文件路径
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()

    # --- 设置默认段落样式 ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

    if not segments:
        # 空转录：生成占位内容
        title = doc.add_paragraph()
        title_run = title.add_run("⚠️ 未检测到语音内容")
        title_run.font.name = "Calibri"
        title_run.font.size = Pt(14)
        title_run.bold = True

        doc.add_paragraph(
            "该音频文件转录为空，可能原因：\n"
            "• 音频为纯音乐或纯器乐，不含人声\n"
            "• 音频全程为静音或环境噪音\n"
            "• 语言与模型支持的语言不匹配"
        )

        footer = doc.add_paragraph()
        footer_run = footer.add_run(
            "此文件由 Video-Transcriber 自动生成"
        )
        footer_run.font.name = "Calibri"
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.save(output_path)
        return output_path

    # --- 逐段写入 ---
    for seg in segments:
        text = seg.get("text", "").strip()
        if not text:
            continue

        para = doc.add_paragraph()

        if with_timestamps:
            # 时间戳 run：等宽字体 + 深灰色
            ts = _format_timestamp(seg["start"])
            ts_run = para.add_run(f"{ts} ")
            ts_run.font.name = "Consolas"
            ts_run.font.size = Pt(10)
            ts_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # 正文 run
            text_run = para.add_run(text)
            text_run.font.name = "Calibri"
            text_run.font.size = Pt(11)
        else:
            text_run = para.add_run(text)
            text_run.font.name = "Calibri"
            text_run.font.size = Pt(11)

    doc.save(output_path)
    return output_path


# ============================================================
# 便捷入口：根据配置生成输出文件
# ============================================================

def save_transcript(
    transcript_result: Dict[str, Any],
    input_file_path: str,
    config: Any,
) -> List[str]:
    """根据配置将转录结果保存为一个或多个格式化文件。

    根据 config.output_format 决定输出格式：
    - 'md'：仅输出 Markdown 文件
    - 'docx'：仅输出 Word 文件
    - 'both'：同时输出 .md 和 .docx 两份文件

    如果转录结果为空（无语音内容），仍然生成输出文件，
    文件中包含"未检测到语音内容"的提示信息。

    输出文件保存在 config.output_output_dir 指定的目录中，
    文件名取自输入文件的主名（去掉扩展名），扩展名相应变化。

    Args:
        transcript_result: WhisperTranscriber.transcribe() 返回的结果字典，
                           包含 'segments' 键
        input_file_path: 原始输入文件路径（用于确定输出文件名）
        config: AppConfig 配置对象（来自 config.py）

    Returns:
        生成的文件路径列表

    示例:
        >>> result = transcriber.transcribe("meeting.wav")
        >>> paths = save_transcript(result, "video.mp4", config)
        >>> print(paths)
        ['./transcripts/video.md', './transcripts/video.docx']
    """
    segments = transcript_result.get("segments", [])

    # 确定输出文件主名（去掉原始扩展名）
    stem = Path(input_file_path).stem
    output_dir = config.output_output_dir
    fmt = config.output_format
    with_ts = config.output_with_timestamps

    output_files: List[str] = []

    if fmt in ("md", "both"):
        md_path = os.path.join(output_dir, f"{stem}.md")
        save_markdown(segments, md_path, with_timestamps=with_ts)
        output_files.append(md_path)

    if fmt in ("docx", "both"):
        docx_path = os.path.join(output_dir, f"{stem}.docx")
        save_docx(segments, docx_path, with_timestamps=with_ts)
        output_files.append(docx_path)

    return output_files
